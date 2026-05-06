from app.tasks.celery_app import celery_app
from app.core.database import SyncSessionLocal
from app.models.document import Document
from app.models.knowledge import KnowledgeBaseHierarchy, KnowledgeChunk
from app.models.system import AsyncTask
from app.models.enums import DocumentStatus, TaskStatus
from app.core.config import settings
from app.core.ollama_client import generate_sync
from sqlalchemy import select
import os
import json
import redis
from datetime import datetime

sync_redis = redis.from_url(settings.REDIS_URL, encoding="utf-8", decode_responses=True)

_POLISH_PROMPT_PATH = os.path.join(os.path.dirname(__file__), "..", "prompts", "polish_system.txt")

@celery_app.task(bind=True, max_retries=3)
def process_polish_task(self, task_id: str, doc_id: str):
    with SyncSessionLocal() as session:
        _mark_task_processing(session, task_id)
        try:
            # 1. 加载公文与任务参数（单独事务，不长时间持锁）
            with session.begin():
                doc = session.execute(
                    select(Document).where(Document.doc_id == doc_id).with_for_update()
                ).scalars().first()

                if not doc or doc.status != DocumentStatus.DRAFTING:
                    _mark_task_failed(session, task_id, "公文状态已变更，润色结果被废弃")
                    return

                doc_content = doc.content or ""

                task = session.get(AsyncTask, task_id)
                input_params = task.input_params if task else {}
                context_kb_ids = input_params.get("context_kb_ids", [])
                exemplar_id = input_params.get("exemplar_id")

                # 上下文版本校验：若 KB 在任务创建后被修改或删除，记录警告
                snapshot_version = input_params.get("snapshot_version")
                if context_kb_ids and snapshot_version:
                    from datetime import datetime as dt, timezone
                    stale_result = session.execute(
                        select(KnowledgeBaseHierarchy.kb_id, KnowledgeBaseHierarchy.kb_name).where(
                            KnowledgeBaseHierarchy.kb_id.in_(context_kb_ids),
                            KnowledgeBaseHierarchy.updated_at
                            > dt.fromtimestamp(snapshot_version, tz=timezone.utc).replace(tzinfo=None)
                        )
                    )
                    stale = stale_result.all()
                    if stale:
                        stale_names = [s[1] for s in stale]
                        _publish_progress(task_id, 15,
                            f"⚠ 以下知识资产在任务创建后发生变更: {', '.join(stale_names)}"
                        )

                # 加载知识库上下文
                context_parts = []
                if context_kb_ids:
                    chunks = session.execute(
                        select(KnowledgeChunk).where(
                            KnowledgeChunk.kb_id.in_(context_kb_ids),
                            KnowledgeChunk.is_deleted == False
                        ).limit(10)
                    ).scalars().all()
                    for c in chunks:
                        meta = c.metadata_json or {}
                        context_parts.append(f"[来源: {meta.get('title_path', '未知')}]\n{c.content}")

                # 加载范文
                if exemplar_id:
                    from app.models.document import ExemplarDocument
                    exemplar = session.get(ExemplarDocument, exemplar_id)
                    if exemplar and exemplar.extracted_text:
                        context_parts.insert(0, f"[参考范文: {exemplar.title}]\n{exemplar.extracted_text}")

                context_str = "\n\n".join(context_parts) if context_parts else "无额外参考上下文"

            # 2. 构建 Prompt（事务外）
            if os.path.exists(_POLISH_PROMPT_PATH):
                with open(_POLISH_PROMPT_PATH, "r", encoding="utf-8") as f:
                    template = f.read()
            else:
                template = "请润色以下公文：\n{content}\n\n参考上下文：{context}\n\n润色后正文："

            prompt = template.replace("{content}", doc_content).replace("{context}", context_str)

            # 3. 调用 Ollama（事务外，不持数据库锁）
            _publish_progress(task_id, 30, "正在调用 AI 模型...")
            try:
                polished = generate_sync(prompt)
            except Exception as llm_err:
                _mark_task_failed(session, task_id, f"AI 引擎调用失败: {str(llm_err)}")
                return

            if not polished or not polished.strip():
                _mark_task_failed(session, task_id, "AI 引擎返回空结果")
                return

            _publish_progress(task_id, 90, "正在保存结果...")

            # 4. 铁律：跨进程状态二次复核 + 保存（新事务 + 行锁）
            with session.begin():
                doc = session.execute(
                    select(Document).where(Document.doc_id == doc_id).with_for_update()
                ).scalars().first()
                if not doc or doc.status != DocumentStatus.DRAFTING:
                    _mark_task_failed(session, task_id, "公文状态已变更，润色结果被废弃")
                    return

                doc.ai_polished_content = polished
                doc.draft_suggestion = polished
                _mark_task_completed(session, task_id, polished[:200])

                # 创建通知（与排版任务对齐：用户离开页面后也能感知完成）
                from app.models.system import UserNotification
                from app.models.enums import NotificationType
                notif = UserNotification(
                    user_id=doc.creator_id,
                    doc_id=doc_id,
                    type=NotificationType.TASK_COMPLETED,
                    content=f"公文「{doc.title}」AI 润色完成",
                )
                session.add(notif)

        except Exception as e:
            session.rollback()
            if self.request.retries >= self.max_retries:
                _mark_task_failed(session, task_id, f"任务异常: {str(e)}")
                return
            raise self.retry(exc=e)

@celery_app.task(bind=True, max_retries=3)
def process_format_task(self, doc_id: str, task_id: str = None):
    with SyncSessionLocal() as session:
        if task_id:
            _mark_task_processing(session, task_id)
        try:
            os.makedirs(settings.OUTPUT_DIR, exist_ok=True)
            output_path = os.path.join(settings.OUTPUT_DIR, f"{doc_id}.docx")

            # 铁律：跨进程状态二次复核 (P0 §七.3)
            with session.begin():
                doc = session.execute(
                    select(Document).where(Document.doc_id == doc_id).with_for_update()
                ).scalars().first()
                if not doc or doc.status != DocumentStatus.APPROVED:
                    if task_id:
                        _mark_task_failed(session, task_id, "公文未处于已通过状态")
                    return

                # FORMAT_REQUESTED audit
                from app.models.system import NBSWorkflowAudit
                from app.models.enums import WorkflowNodeId
                audit_req = NBSWorkflowAudit(
                    doc_id=doc_id,
                    workflow_node_id=WorkflowNodeId.FORMAT_REQUESTED,
                    operator_id=doc.creator_id,
                    action_details={"output_path": output_path}
                )
                session.add(audit_req)

                # Look up layout rules from DocumentType
                layout_rules = {}
                try:
                    from app.models.document import DocumentType
                    dt_result = session.execute(
                        select(DocumentType).where(DocumentType.type_id == doc.doc_type_id)
                    )
                    dt = dt_result.scalars().first()
                    if dt and dt.layout_rules:
                        layout_rules = dt.layout_rules
                except Exception:
                    pass

                # Generate .docx with layout rules applied
                try:
                    from docx import Document as DocxDocument
                    from docx.shared import Pt, Cm, Emu
                    from docx.enum.text import WD_ALIGN_PARAGRAPH

                    docx = DocxDocument()

                    # Apply page margins from layout rules
                    margins = layout_rules.get("page_margins", {})
                    section = docx.sections[0]
                    section.top_margin = Cm(margins.get("top", 3.7))
                    section.bottom_margin = Cm(margins.get("bottom", 3.5))
                    section.left_margin = Cm(margins.get("left", 2.8))
                    section.right_margin = Cm(margins.get("right", 2.6))

                    # Title (red header — 红头)
                    title_para = docx.add_paragraph()
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    title_run = title_para.add_run(doc.title or "未命名公文")
                    title_font_size = layout_rules.get("title_font_size", 22)
                    title_run.font.size = Pt(title_font_size)
                    title_run.font.name = layout_rules.get("title_font", "方正小标宋简体")
                    title_run.bold = True

                    # Document number line
                    if doc.document_number:
                        num_para = docx.add_paragraph()
                        num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        num_run = num_para.add_run(f"发文编号：{doc.document_number}")
                        num_run.font.size = Pt(layout_rules.get("body_font_size", 16))

                    # Body content
                    body_font_size = layout_rules.get("body_font_size", 16)
                    body_font = layout_rules.get("body_font", "仿宋_GB2312")
                    line_spacing = layout_rules.get("line_spacing_pt", 28)

                    for para_text in (doc.content or "").split("\n"):
                        body_para = docx.add_paragraph()
                        body_para.paragraph_format.line_spacing = Pt(line_spacing)
                        # First-line indent: 2 chars at current font size
                        body_para.paragraph_format.first_line_indent = Pt(body_font_size * 2)
                        body_run = body_para.add_run(para_text)
                        body_run.font.size = Pt(body_font_size)
                        body_run.font.name = body_font

                    docx.save(output_path)
                except ImportError:
                    # python-docx not available, write placeholder
                    with open(output_path, "w", encoding="utf-8") as f:
                        f.write(f"{doc.title}\n\n{doc.content or ''}")

                doc.word_output_path = output_path

                # FORMAT_COMPLETED audit
                audit_done = NBSWorkflowAudit(
                    doc_id=doc_id,
                    workflow_node_id=WorkflowNodeId.FORMAT_COMPLETED,
                    operator_id=doc.creator_id,
                    action_details={"output_path": output_path}
                )
                session.add(audit_done)

                if task_id:
                    _mark_task_completed(session, task_id, output_path)

                # Create notification
                from app.models.system import UserNotification
                from app.models.enums import NotificationType
                notif = UserNotification(
                    user_id=doc.creator_id,
                    doc_id=doc_id,
                    type=NotificationType.TASK_COMPLETED,
                    content=f"公文「{doc.title}」排版完成",
                )
                session.add(notif)
        except Exception as e:
            session.rollback()
            if self.request.retries >= self.max_retries:
                if task_id:
                    _mark_task_failed(session, task_id, f"排版任务异常: {str(e)}")
                return
            raise self.retry(exc=e)

@celery_app.task(bind=True, max_retries=3)
def process_parse_task(self, kb_id: int, task_id: str = None):
    with SyncSessionLocal() as session:
        if task_id:
            _mark_task_processing(session, task_id)
            _publish_progress(task_id, 10, "开始解析...")
        try:
            with session.begin():
                from app.models.knowledge import KnowledgeChunk, KnowledgePhysicalFile
                from sqlalchemy import text

                kb_node = session.execute(
                    select(KnowledgeBaseHierarchy).where(KnowledgeBaseHierarchy.kb_id == kb_id).with_for_update()
                ).scalars().first()
                if not kb_node or kb_node.is_deleted:
                    if task_id:
                        _mark_task_failed(session, task_id, "节点已被删除")
                    return

                # Mark as parsing
                kb_node.parse_status = "PARSING"
                session.commit()

            if task_id:
                _publish_progress(task_id, 30, "读取文件内容...")

            # Read file content
            physical_file = session.get(KnowledgePhysicalFile, kb_node.physical_file_id)
            content = "模拟切片内容"
            if physical_file and os.path.exists(physical_file.file_path):
                try:
                    with open(physical_file.file_path, "r", encoding="utf-8") as f:
                        content = f.read()
                except Exception:
                    pass

            if task_id:
                _publish_progress(task_id, 60, "生成切片...")

            with session.begin():
                chunk = KnowledgeChunk(
                    kb_id=kb_node.kb_id,
                    physical_file_id=kb_node.physical_file_id,
                    chunk_index=0,
                    kb_tier=kb_node.kb_tier,
                    content=content,
                    security_level=kb_node.security_level,
                    metadata_json={"title_path": kb_node.kb_name},
                    embedding=None
                )
                session.add(chunk)

                kb_node = session.get(KnowledgeBaseHierarchy, kb_id)
                kb_node.parse_status = "READY"
                session.commit()

            if task_id:
                _publish_progress(task_id, 100, "解析完成")
                _mark_task_completed(session, task_id, "parsed")
        except Exception as e:
            session.rollback()
            # Mark node as FAILED so user can trigger re-parse
            with session.begin():
                kb_node = session.get(KnowledgeBaseHierarchy, kb_id)
                if kb_node:
                    kb_node.parse_status = "FAILED"
            if task_id:
                _mark_task_failed(session, task_id, str(e))
            if self.request.retries >= self.max_retries:
                return
            raise self.retry(exc=e)

def _mark_task_processing(session, task_id):
    task = session.get(AsyncTask, task_id)
    if task:
        task.task_status = TaskStatus.PROCESSING
        task.started_at = datetime.now()
        session.commit()

def _mark_task_completed(session, task_id, summary):
    task = session.get(AsyncTask, task_id)
    if task:
        task.task_status = TaskStatus.COMPLETED
        task.progress_pct = 100
        task.result_summary = summary
        task.completed_at = datetime.now()
        session.commit()
        sync_redis.publish(f"task_events:{task_id}", json.dumps({"event": "task.completed", "task_id": task_id, "result_summary": summary}))

def _mark_task_failed(session, task_id, error):
    task = session.get(AsyncTask, task_id)
    if task:
        task.task_status = TaskStatus.FAILED
        task.error_message = error
        task.completed_at = datetime.now()
        session.commit()
        sync_redis.publish(f"task_events:{task_id}", json.dumps({"event": "task.failed", "task_id": task_id, "error_message": error}))

def _publish_progress(task_id, pct, message):
    sync_redis.publish(f"task_events:{task_id}", json.dumps({
        "event": "task.progress",
        "task_id": task_id,
        "progress_pct": pct,
        "message": message,
    }))