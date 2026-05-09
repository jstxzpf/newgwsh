"""
GB/T 9704-2023 国标模板化排版引擎
按 DocumentType.layout_rules.template 逐段渲染 .docx
占位符格式：【字段名】
"""
from docx import Document as DocxDocument
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import re
import os

_PLACEHOLDER_RE = re.compile(r"【(.+?)】")

DEFAULT_MARGINS = {"top": 3.7, "bottom": 3.5, "left": 2.8, "right": 2.6}
DEFAULT_FONT = "仿宋_GB2312"
DEFAULT_SIZE = 16
DEFAULT_LINE_SPACING = Pt(28)

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def _resolve(text: str, data: dict) -> str:
    def replacer(m):
        key = m.group(1)
        val = data.get(key)
        if val:
            return str(val)
        return ""
    return _PLACEHOLDER_RE.sub(replacer, text)


def _build_data(doc, layout_rules: dict) -> dict:
    dept_name = layout_rules.get("dept_name", "")
    return {
        "发文机关名称": dept_name,
        "发文字号": getattr(doc, "document_number", None) or "",
        "签发人": "",
        "标题": getattr(doc, "title", None) or "未命名公文",
        "主送机关": getattr(doc, "recipient", None) or "",
        "正文": (getattr(doc, "ai_polished_content", None) or getattr(doc, "content", None) or ""),
        "附件说明": "",
        "发文机关署名": dept_name,
        "成文日期": _format_date(getattr(doc, "issued_at", None) or getattr(doc, "approved_at", None)),
        "抄送机关": getattr(doc, "cc_list", None) or "",
        "印发日期": _format_date(datetime.now()),
        "联系人": layout_rules.get("contact", ""),
        "联系电话": layout_rules.get("contact_phone", ""),
    }


def _format_date(dt) -> str:
    if dt is None:
        return datetime.now().strftime("%Y年%m月%d日")
    if isinstance(dt, str):
        return dt
    return dt.strftime("%Y年%m月%d日")


_HAS_CONTENT_CONDITIONS = {
    "has_attachments": lambda data: bool(data.get("附件说明", "").strip()),
    "has_cc": lambda data: bool(data.get("抄送机关", "").strip()),
    "has_recipient": lambda data: bool(data.get("主送机关", "").strip()),
    "has_doc_number": lambda data: bool(data.get("发文字号", "").strip()),
}


def render(doc, layout_rules: dict, output_path: str):
    """主入口：将公文数据 + 模板渲染为 .docx 文件"""
    data = _build_data(doc, layout_rules)
    template = layout_rules.get("template", [])
    margins = layout_rules.get("page_margins", DEFAULT_MARGINS)

    if not template:
        _render_fallback(doc, layout_rules, output_path)
        return

    docx = DocxDocument()

    section = docx.sections[0]
    section.top_margin = Cm(margins.get("top", 3.7))
    section.bottom_margin = Cm(margins.get("bottom", 3.5))
    section.left_margin = Cm(margins.get("left", 2.8))
    section.right_margin = Cm(margins.get("right", 2.6))

    for item in template:
        _render_section(docx, item, data)

    docx.save(output_path)


def _render_section(docx, item: dict, data: dict):
    section_type = item.get("type", "paragraph")

    # 条件渲染
    cond = item.get("condition")
    if cond:
        checker = _HAS_CONTENT_CONDITIONS.get(cond)
        if checker and not checker(data):
            return
        if not checker:
            resolved = _resolve(item.get("text", ""), data)
            if not resolved.strip():
                return

    if section_type == "separator":
        _render_separator(docx, item)
    elif section_type == "body":
        _render_body(docx, item, data)
    elif section_type == "red_header":
        _render_paragraph(docx, {**item, "align": "center"})
    elif section_type == "ending":
        _render_ending(docx, item, data)
    else:
        _render_paragraph(docx, item)


def _render_paragraph(docx, item: dict):
    text = item.get("text", "")
    font_name = item.get("font", DEFAULT_FONT)
    font_size = item.get("size", DEFAULT_SIZE)
    align = item.get("align", "left")
    bold = item.get("bold", False)
    color = item.get("color")
    space_before = item.get("space_before", 0)
    space_after = item.get("space_after", 0)

    para = docx.add_paragraph()
    para.alignment = ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    if space_before:
        para.paragraph_format.space_before = Pt(space_before)
    if space_after:
        para.paragraph_format.space_after = Pt(space_after)

    if not text:
        return

    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _render_ending(docx, item: dict, data: dict):
    text = _resolve(item.get("text", ""), data)
    if not text.strip():
        return
    _render_paragraph(docx, item)


def _render_body(docx, item: dict, data: dict):
    text = _resolve(item.get("text", ""), data)
    if not text.strip():
        return

    font_name = item.get("font", DEFAULT_FONT)
    font_size = item.get("size", DEFAULT_SIZE)
    line_spacing_val = item.get("line_spacing", 28)
    first_line_indent = item.get("first_line_indent", font_size * 2)

    for para_text in text.split("\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        para = docx.add_paragraph()
        para.paragraph_format.line_spacing = Pt(line_spacing_val)
        para.paragraph_format.first_line_indent = Pt(first_line_indent)
        run = para.add_run(para_text)
        run.font.size = Pt(font_size)
        run.font.name = font_name


def _render_separator(docx, item: dict):
    style = item.get("style", "red_line")
    para = docx.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    if style == "red_line":
        run = para.add_run("─" * 50)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)


def _render_fallback(doc, layout_rules: dict, output_path: str):
    """无模板时的兜底排版（保持向后兼容）"""
    docx = DocxDocument()
    margins = layout_rules.get("page_margins", DEFAULT_MARGINS)
    section = docx.sections[0]
    section.top_margin = Cm(margins["top"])
    section.bottom_margin = Cm(margins["bottom"])
    section.left_margin = Cm(margins["left"])
    section.right_margin = Cm(margins["right"])

    title_para = docx.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(getattr(doc, "title", None) or "未命名公文")
    run.font.size = Pt(22)
    run.font.name = "方正小标宋简体"
    run.bold = True

    body_text = getattr(doc, "ai_polished_content", None) or getattr(doc, "content", None) or ""
    for para_text in body_text.split("\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        para = docx.add_paragraph()
        para.paragraph_format.line_spacing = DEFAULT_LINE_SPACING
        para.paragraph_format.first_line_indent = Pt(32)
        run = para.add_run(para_text)
        run.font.size = Pt(DEFAULT_SIZE)
        run.font.name = DEFAULT_FONT

    docx.save(output_path)
