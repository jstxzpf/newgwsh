import os
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from markdown_it import MarkdownIt

class DocxService:
    def __init__(self):
        self.md_it = MarkdownIt("commonmark")

    def apply_gb_styles(self, doc):
        """应用国标基础样式 (GB/T 9704-2023)"""
        # 设置页边距 (上下 37mm, 左右 28mm 为近似值)
        sections = doc.sections
        for section in sections:
            section.top_margin = Mm(37)
            section.bottom_margin = Mm(35)
            section.left_margin = Mm(28)
            section.right_margin = Mm(26)
            section.page_height = Mm(297)
            section.page_width = Mm(210)

    def set_font(self, run, font_name, size_pt=16):
        """设置中文字体和字号 (三号 = 16pt)"""
        run.font.size = Pt(size_pt)
        run.font.name = font_name
        # 强制设置中文字体
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), font_name)

    def format_document(self, content_md: str, output_path: str, doc_type_rules: dict = None):
        """将 Markdown 转换为符合国标的 Docx (含文种分支路由 P3.3)"""
        # 1. 词库清洗 (实施约束规则 3)
        from app.services.prompt_service import prompt_loader
        blacklist = prompt_loader.get_vocab_blacklist()
        for old, new in blacklist.items():
            content_md = content_md.replace(old, new)

        # 2. 文种路由判断
        required_sections = doc_type_rules.get("required_sections", []) if doc_type_rules else []
        is_basic_mode = len(required_sections) == 0
        
        doc = DocxDocument()
        self.apply_gb_styles(doc)
        
        # 3. 结构校验 (非基础模式 P3.3)
        warnings = []
        if not is_basic_mode:
            # 针对 RESEARCH 与 ECONOMIC_INFO 进行专业结构校验
            # required_sections 示例: ["调研背景", "主要发现", "政策建议"]
            for section in required_sections:
                # 简单匹配: 检查 Markdown 中是否包含该标题或关键词
                if section not in content_md:
                    warnings.append(f"结构缺失警告: 当前文档未检测到【{section}】章节，请核实。")

        # 4. AST 遍历与渲染 (全线代码明确锁定唯一授权使用 markdown-it-py P1.4)
        tokens = self.md_it.parse(content_md)
        
        for i, token in enumerate(tokens):
            if token.type == "heading_open":
                level = int(token.tag[1])
                inline_token = tokens[i+1]
                p = doc.add_paragraph()
                run = p.add_run(inline_token.content)
                
                if level == 1: # ## 一级标题 (黑体, 三号)
                    self.set_font(run, "黑体", 16)
                elif level == 2: # ### 二级标题 (楷体)
                    self.set_font(run, "楷体_GB2312", 16)
                else:
                    self.set_font(run, "仿宋_GB2312", 16)
                
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                p.paragraph_format.line_spacing = Pt(28)
                
            elif token.type == "paragraph_open":
                inline_token = tokens[i+1]
                if inline_token.type == "inline":
                    p = doc.add_paragraph()
                    # 首行缩进 2 字符
                    p.paragraph_format.first_line_indent = Pt(32) 
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    p.paragraph_format.line_spacing = Pt(28)
                    
                    run = p.add_run(inline_token.content)
                    self.set_font(run, "仿宋_GB2312", 16)
                    
        doc.save(output_path)
        return {"file_path": output_path, "warnings": warnings, "mode": "basic" if is_basic_mode else "standard"}

docx_service = DocxService()
